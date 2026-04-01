"""
Script para generar el informe Word descriptivo del proyecto
"Generador de Informes Automatizados (TrackingTime & Azure DevOps)"

Ejecutar: python generar_informe_word.py
Salida:   output/Informe_Descriptivo_Entregable.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda de tabla."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Crea una tabla con formato profesional."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Encabezados
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
                run.font.name = 'Arial'
        set_cell_shading(hdr_cells[i], "1F3864")

    # Datos
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, value in enumerate(row_data):
            row_cells[col_idx].text = str(value)
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Arial'
            if row_idx % 2 == 1:
                set_cell_shading(row_cells[col_idx], "F2F2F2")

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    return table


def create_report():
    doc = Document()

    # ── Estilos globales ──
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(6)

    # ═══════════════════════════════════════════════════════════════════════
    # PORTADA
    # ═══════════════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("INFORME DESCRIPTIVO DEL ENTREGABLE")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 56, 100)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Generador de Informes Automatizados\nTrackingTime & Azure DevOps")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("Logiztik Alliance — Área de Tecnología")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # TABLA DE CONTENIDOS (manual)
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading("Tabla de Contenidos", level=1)
    toc_items = [
        "1. Objetivo del Documento",
        "2. Descripción General del Proyecto",
        "3. Estructura de Carpetas del Proyecto",
        "4. Archivos de Referencia (Carpeta Documentos)",
        "5. Fase 1 — Extracción de Datos desde las APIs",
        "   5.1. Extracción de Azure DevOps",
        "   5.2. Extracción de TrackingTime",
        "6. Archivos CSV Intermedios Generados",
        "7. Fase 2 — Transformación de Datos",
        "   7.1. Transformación de Azure DevOps",
        "   7.2. Transformación de TrackingTime",
        "   7.3. Construcción de Tablas Derivadas",
        "   7.4. Tabla de Presentación Final",
        "8. Informe Excel Final — Estructura del Archivo",
        "   8.1. Pestaña PRESENTACION",
        "   8.2. Pestaña CODIGOS PROYECTOS",
        "   8.3. Pestaña DETALLE - AZURE",
        "   8.4. Pestaña DETALLE - TRACKINGTIME",
        "9. Flujo de Datos Completo (Resumen)",
        "10. Configuración y Variables de Entorno",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 1. OBJETIVO
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading("1. Objetivo del Documento", level=1)
    doc.add_paragraph(
        "Este documento tiene como propósito describir de manera exhaustiva el entregable "
        "correspondiente al proyecto de automatización de informes de horas trabajadas. "
        "Se detalla la estructura de carpetas, el contenido de cada archivo, las fuentes de datos "
        "(APIs de Azure DevOps y TrackingTime), las transformaciones aplicadas a los datos crudos "
        "y el significado de cada columna presente en el informe Excel final."
    )
    doc.add_paragraph(
        "El objetivo es que cualquier persona que reciba este entregable pueda comprender qué contiene "
        "la carpeta, por qué existe cada archivo, y cómo interpretar la información presentada en el reporte."
    )

    # ═══════════════════════════════════════════════════════════════════════
    # 2. DESCRIPCIÓN GENERAL
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading("2. Descripción General del Proyecto", level=1)
    doc.add_paragraph(
        "El proyecto «Generador de Informes Automatizados» es una aplicación de escritorio desarrollada "
        "en Python que automatiza la extracción, procesamiento y consolidación de datos provenientes de "
        "dos plataformas:"
    )
    items = [
        ("Azure DevOps", "Gestión de tareas y proyectos del equipo de desarrollo. Se extraen los Work Items "
         "(tareas, bugs, historias de usuario) con las horas completadas (Completed Work) por cada miembro del equipo."),
        ("TrackingTime", "Plataforma de registro de tiempo. Se extraen los eventos de tiempo (horas trabajadas) "
         "por cada usuario activo en el rango de fechas configurado.")
    ]
    for name, desc in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph(
        "El programa unifica ambas fuentes y genera un informe Excel consolidado que muestra "
        "las horas trabajadas por cada persona, discriminadas por proyecto, con el código de proyecto "
        "correspondiente y el porcentaje de tiempo asignado a cada proyecto respecto del total de horas "
        "trabajadas por cada individuo."
    )
    doc.add_paragraph(
        "El proceso se ejecuta en dos fases:"
    )
    phases = [
        ("Fase 1 — Extracción:", "Se conecta a las APIs de Azure DevOps y TrackingTime, descarga los datos "
         "y los almacena como archivos CSV intermedios en la carpeta output/."),
        ("Fase 2 — Transformación y Reporte:", "Lee los CSV intermedios, aplica transformaciones de limpieza "
         "y normalización, cruza los datos con los archivos de referencia (CodigosProyectos.csv y Equipo.csv), "
         "construye tablas derivadas y genera el archivo Excel final.")
    ]
    for name, desc in phases:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(name + " ")
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 3. ESTRUCTURA DE CARPETAS
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading("3. Estructura de Carpetas del Proyecto", level=1)
    doc.add_paragraph("A continuación se describe la estructura completa de la carpeta del proyecto:")

    tree = """project/
├── .env                          ← Variables de entorno (tokens, fechas)
├── .env.example                  ← Plantilla de ejemplo del .env
├── .gitignore                    ← Archivos excluidos del control de versiones
├── README.md                     ← Documentación de uso del programa
├── requirements.txt              ← Dependencias de Python
├── config.py                     ← Lectura y validación de configuración
├── main.py                       ← Punto de entrada principal del programa
│
├── Documentos/                   ← Archivos CSV de referencia (catálogos)
│   ├── CodigosProyectos.csv      ← Catálogo maestro de proyectos
│   └── Equipo.csv                ← Catálogo maestro de personal
│
├── azure_devops/                 ← Módulo de extracción de Azure DevOps
│   ├── __init__.py
│   ├── client.py                 ← Cliente HTTP para la API de Azure DevOps
│   └── exporter.py               ← Lógica de exportación y aplanamiento de datos
│
├── trackingtime/                 ← Módulo de extracción de TrackingTime
│   ├── __init__.py
│   ├── client.py                 ← Cliente HTTP para la API de TrackingTime
│   └── exporter.py               ← Lógica de exportación y unificación de CSV
│
├── processing/                   ← Módulo de transformación y generación del reporte
│   ├── __init__.py               ← Exporta las funciones públicas del módulo
│   ├── loader.py                 ← Carga y parseo de los CSV (incluye tipado)
│   ├── transformations.py        ← Limpieza, normalización y tablas derivadas
│   └── excel_writer.py           ← Generación del archivo Excel final
│
└── output/                       ← Carpeta de salida (archivos generados)
    ├── azure_devops_unified_*.csv   ← CSV crudo de Azure DevOps
    ├── trackingtime_unified_*.csv   ← CSV crudo de TrackingTime
    └── Reporte_*.xlsx               ← Informe Excel final"""

    p = doc.add_paragraph()
    run = p.add_run(tree)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)

    doc.add_paragraph()

    folder_desc = [
        ("project/", "Carpeta raíz del proyecto. Contiene todos los archivos fuente, configuración y módulos."),
        ("Documentos/", "Contiene los archivos CSV de referencia que actúan como catálogos maestros. "
         "Estos archivos deben mantenerse actualizados manualmente cuando se agreguen nuevos proyectos o personas al equipo."),
        ("azure_devops/", "Módulo Python responsable de la conexión y extracción de datos desde la API REST de Azure DevOps. "
         "Obtiene los Work Items de un query predefinido."),
        ("trackingtime/", "Módulo Python responsable de la conexión y extracción de datos desde la API REST de TrackingTime. "
         "Exporta los eventos de tiempo de cada usuario activo."),
        ("processing/", "Módulo Python que contiene toda la lógica de la Fase 2: carga de datos, transformaciones, "
         "construcción de tablas derivadas y generación del archivo Excel."),
        ("output/", "Carpeta donde se almacenan todos los archivos generados por el programa: "
         "los CSV intermedios y el reporte Excel final."),
    ]
    for folder, desc in folder_desc:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(folder + " ")
        run.bold = True
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        p.add_run(desc)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 4. ARCHIVOS DE REFERENCIA
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading("4. Archivos de Referencia (Carpeta Documentos/)", level=1)
    doc.add_paragraph(
        "La carpeta Documentos/ contiene dos archivos CSV que actúan como catálogos maestros. "
        "El programa los utiliza para cruzar información y enriquecer los datos extraídos de las APIs."
    )

    # 4.1 CodigosProyectos
    doc.add_heading("4.1. CodigosProyectos.csv", level=2)
    doc.add_paragraph(
        "Este archivo es el catálogo maestro de proyectos. Define la relación entre los nombres "
        "de proyecto en Azure DevOps, los nombres en TrackingTime, y el código interno de proyecto. "
        "También establece el porcentaje de distribución de costos por sucursal (país)."
    )
    doc.add_paragraph("Separador: punto y coma (;)")
    doc.add_paragraph("Columnas:")

    codigos_cols = [
        ("Codigo Proyecto", "Código interno único del proyecto (ej: LAG-001-CL360). "
         "Sigue la convención LAG-NNN-SIGLA donde LAG identifica a Logiztik Alliance Group, "
         "NNN es un consecutivo y SIGLA es una abreviatura del proyecto."),
        ("Proyecto", "Nombre oficial del proyecto tal como se presenta en los reportes."),
        ("Nombre Azure", "Nombre exacto del proyecto en Azure DevOps (campo Team Project). "
         "Se usa para cruzar los datos de Azure con este catálogo."),
        ("Nombre Tracking", "Nombre exacto del proyecto en TrackingTime (campo Project). "
         "Se usa para cruzar los datos de TrackingTime con este catálogo."),
        ("Fecha Registro", "Fecha en que el proyecto fue registrado en el catálogo."),
        ("Descripción corta", "Breve descripción del alcance y objetivo del proyecto."),
        ("Area Beneficiada", "Áreas de la empresa que se benefician de este proyecto."),
        ("ECU", "Porcentaje de distribución del costo hacia la sucursal de Ecuador (ej: 25% → 0.25 internamente)."),
        ("COL", "Porcentaje de distribución del costo hacia la sucursal de Colombia."),
        ("USA", "Porcentaje de distribución del costo hacia la sucursal de Estados Unidos (Miami)."),
        ("NL", "Porcentaje de distribución del costo hacia la sucursal de Países Bajos."),
    ]
    add_styled_table(doc, ["Columna", "Descripción"], codigos_cols, col_widths=[4, 14])

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Nota importante: ")
    run.bold = True
    p.add_run("Un mismo Codigo Proyecto puede aparecer múltiples veces si el mismo código agrupa "
              "varios proyectos de Azure/TrackingTime distintos (por ejemplo, LAG-010-OTROS agrupa "
              "Administración, AC Architecture, AC QA, Base de Conocimiento, etc.).")

    # 4.2 Equipo
    doc.add_heading("4.2. Equipo.csv", level=2)
    doc.add_paragraph(
        "Este archivo es el catálogo maestro del equipo de trabajo. Permite mapear los nombres "
        "que aparecen en Azure DevOps y TrackingTime con los datos formales de cada persona."
    )
    doc.add_paragraph("Separador: punto y coma (;)")
    doc.add_paragraph("Columnas:")

    equipo_cols = [
        ("EMPRESA", "Empresa a la que pertenece el colaborador: LOGIZTIK o NEWDATA (empresa externa/outsourcing)."),
        ("NOMBRE", "Nombre completo formal del colaborador. Este es el nombre que aparecerá en el reporte final."),
        ("Nombre Azure", "Nombre corto del colaborador tal como aparece en Azure DevOps y TrackingTime. "
         "Se utiliza como llave de cruce entre las fuentes de datos."),
        ("CÉDULA DE IDENTIDAD", "Número de documento de identidad del colaborador."),
    ]
    add_styled_table(doc, ["Columna", "Descripción"], equipo_cols, col_widths=[4, 14])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 5. FASE 1 — EXTRACCIÓN
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading("5. Fase 1 — Extracción de Datos desde las APIs", level=1)
    doc.add_paragraph(
        "La primera fase del programa se conecta a las APIs externas y descarga los datos crudos. "
        "Cada fuente de datos tiene su propio módulo con un cliente HTTP y un exportador."
    )

    # 5.1 Azure DevOps
    doc.add_heading("5.1. Extracción de Azure DevOps", level=2)
    doc.add_paragraph("Fuente: API REST de Azure DevOps (versión 7.0)")
    doc.add_paragraph("Módulo: azure_devops/client.py + azure_devops/exporter.py")
    doc.add_paragraph()

    doc.add_heading("Proceso de extracción paso a paso:", level=3)
    steps = [
        "Autenticación: Se utiliza un Personal Access Token (PAT) codificado en Base64 como header Authorization.",
        "Obtener URL del Query: Se consulta el endpoint de queries de Azure DevOps usando el ADO_QUERY_ID configurado. "
        "Se extrae la URL WIQL (Work Item Query Language) del campo _links.wiql.href de la respuesta.",
        "Ejecutar el Query WIQL: Se ejecuta la URL WIQL obtenida, que retorna una lista de IDs de Work Items "
        "que coinciden con los filtros definidos en el query de Azure DevOps.",
        "Obtener detalle de Work Items en lotes: Los IDs se procesan en lotes de 200 elementos "
        "(limitación de la API). Para cada lote, se hace un POST al endpoint workitemsbatch "
        "solicitando campos específicos.",
        "Aplanar la respuesta: Cada Work Item viene como un objeto JSON con un diccionario fields. "
        "Se extrae cada campo y se mapea a una columna plana del CSV.",
        "Formatear fechas: Las fechas se convierten del formato ISO 8601 a formato US (M/D/YYYY h:mm:ss AM/PM).",
        "Guardar CSV: El DataFrame resultante se guarda como azure_devops_unified_YYYYMMDD_YYYYMMDD.csv en la carpeta output/.",
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. ", style='List Number')
        parts = step.split(": ", 1)
        run = p.add_run(parts[0] + ": ")
        run.bold = True
        if len(parts) > 1:
            p.add_run(parts[1])

    doc.add_paragraph()
    doc.add_heading("Campos extraídos de la API de Azure DevOps:", level=3)

    azure_fields = [
        ("System.Id", "ID", "Identificador numérico único del Work Item en Azure DevOps."),
        ("System.State", "State", "Estado actual del Work Item (ej: New, Active, Closed, Resolved)."),
        ("System.AssignedTo", "Assigned To", "Persona asignada al Work Item. Viene como objeto JSON con displayName y uniqueName; "
         "se concatenan como 'Nombre <email>'."),
        ("Microsoft.VSTS.Scheduling.CompletedWork", "Completed Work", "Horas completadas registradas en el Work Item. "
         "Este es el dato central para el cálculo de horas."),
        ("System.TeamProject", "Team Project", "Nombre del proyecto en Azure DevOps al que pertenece el Work Item. "
         "Se usa como llave para cruzar con CodigosProyectos.csv."),
        ("System.WorkItemType", "Work Item Type", "Tipo de elemento: Task, Bug, User Story, etc."),
        ("System.IterationPath", "Iteration Path", "Ruta de la iteración/sprint al que pertenece el Work Item."),
        ("System.Title", "Title", "Título descriptivo del Work Item."),
        ("System.CreatedDate", "Created Date", "Fecha y hora de creación del Work Item."),
        ("Microsoft.VSTS.Common.ClosedDate", "Closed Date", "Fecha y hora en que se cerró el Work Item."),
        ("System.ChangedDate", "Changed Date", "Fecha y hora de la última modificación."),
    ]
    add_styled_table(doc,
                     ["Campo API (original)", "Columna CSV", "Descripción"],
                     azure_fields,
                     col_widths=[5.5, 3, 9.5])

    doc.add_page_break()

    # 5.2 TrackingTime
    doc.add_heading("5.2. Extracción de TrackingTime", level=2)
    doc.add_paragraph("Fuente: API REST de TrackingTime (versión 4)")
    doc.add_paragraph("Módulo: trackingtime/client.py + trackingtime/exporter.py")
    doc.add_paragraph()

    doc.add_heading("Proceso de extracción paso a paso:", level=3)
    steps_tt = [
        "Autenticación: Se utiliza un Token de aplicación (App Password) codificado en Base64 como header Authorization. "
        "No se usa la contraseña habitual de la cuenta.",
        "Obtener lista de usuarios: Se consulta el endpoint /users para obtener todos los usuarios registrados.",
        "Filtrar usuarios activos: Se eliminan los usuarios marcados como archivados (is_archived = true).",
        "Exportar eventos por usuario: Para cada usuario activo, se llama al endpoint /events/export "
        "con los parámetros from (TT_DATE_FROM), to (TT_DATE_TO), type=user, id=<user_id>. "
        "La respuesta es un CSV en texto plano.",
        "Parsear CSV de cada usuario: Cada respuesta CSV se parsea con pandas. Se eliminan filas completamente vacías.",
        "Unificar: Todos los DataFrames individuales se concatenan en uno solo.",
        "Guardar CSV: El DataFrame unificado se guarda como trackingtime_unified_YYYYMMDD_YYYYMMDD.csv en la carpeta output/.",
    ]
    for i, step in enumerate(steps_tt, 1):
        p = doc.add_paragraph(f"{i}. ", style='List Number')
        parts = step.split(": ", 1)
        run = p.add_run(parts[0] + ": ")
        run.bold = True
        if len(parts) > 1:
            p.add_run(parts[1])

    doc.add_paragraph()
    doc.add_heading("Campos del CSV exportado por TrackingTime:", level=3)
    doc.add_paragraph(
        "La API de TrackingTime retorna directamente un CSV con los siguientes campos principales:"
    )
    tt_fields = [
        ("Project", "Nombre del proyecto en TrackingTime al que se le registró el tiempo."),
        ("User", "Nombre del usuario que registró el tiempo."),
        ("Hours", "Cantidad de horas registradas en el evento de tiempo. Puede usar coma decimal (europeo).",),
        ("Service", "Tipo de servicio asociado al evento. Se elimina en la transformación posterior."),
        ("Client", "Cliente asociado al evento. Se elimina en la transformación posterior."),
    ]
    add_styled_table(doc,
                     ["Columna", "Descripción"],
                     tt_fields,
                     col_widths=[4, 14])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 6. CSV INTERMEDIOS
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading("6. Archivos CSV Intermedios Generados", level=1)
    doc.add_paragraph(
        "Al finalizar la Fase 1, se generan dos archivos CSV en la carpeta output/. "
        "El sufijo de fecha refleja el rango configurado en el .env (TT_DATE_FROM y TT_DATE_TO)."
    )

    csv_files = [
        ("azure_devops_unified_YYYYMMDD_YYYYMMDD.csv",
         "Contiene todos los Work Items extraídos de Azure DevOps con los 11 campos descritos en la sección 5.1. "
         "Un registro por cada Work Item encontrado por el query."),
        ("trackingtime_unified_YYYYMMDD_YYYYMMDD.csv",
         "Contiene todos los eventos de tiempo de TrackingTime de todos los usuarios activos, "
         "unificados en un solo archivo. Un registro por cada evento de tiempo."),
    ]
    for name, desc in csv_files:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(name + ": ")
        run.bold = True
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        p.add_run(desc)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Estos archivos son la \"materia prima\" de la Fase 2. ")
    run.italic = True
    p.add_run("Si solo se requiere la información cruda sin transformar, estos CSV se pueden abrir directamente en Excel.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 7. FASE 2 — TRANSFORMACIÓN
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading("7. Fase 2 — Transformación de Datos", level=1)
    doc.add_paragraph(
        "La segunda fase toma los CSV intermedios y los archivos de referencia, aplica transformaciones "
        "de limpieza y normalización, construye tablas derivadas (agregaciones) y genera el reporte Excel final."
    )

    # 7.1 Transform Azure
    doc.add_heading("7.1. Transformación de Azure DevOps", level=2)
    doc.add_paragraph("Archivo: processing/transformations.py → función transform_azure_devops()")
    doc.add_paragraph("Transformaciones aplicadas:")

    azure_transforms = [
        ("Limpieza del campo Assigned To",
         "El campo original viene en formato \"Nombre Apellido Empresa <email@dominio.com>\". "
         "Se extrae solo la parte antes del carácter '<'. Luego se eliminan los sufijos de empresa: "
         "\"Logiztik Alliance\", \"New Data (Externo)\" y \"New Data\". El resultado se guarda en una nueva "
         "columna llamada Assigned To Clean. La columna original Assigned To se elimina."),
        ("Filtrado de filas sin asignación",
         "Se eliminan las filas donde Assigned To Clean está vacío o es nulo, "
         "ya que no se pueden asignar horas a una persona sin nombre."),
    ]
    for name, desc in azure_transforms:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(name + ": ")
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph()
    doc.add_paragraph("Ejemplo de la transformación Assigned To:")
    transform_example = [
        ("Andy Briones Logiztik Alliance <andy@lag.com>", "Andy Briones"),
        ("Geovanny Bastidas New Data (Externo) <geo@nd.com>", "Geovanny Bastidas"),
        ("Jonathan Merino New Data <jon@nd.com>", "Jonathan Merino"),
    ]
    add_styled_table(doc, ["Valor original (Assigned To)", "Valor limpio (Assigned To Clean)"],
                     transform_example, col_widths=[10, 8])

    # 7.2 Transform TrackingTime
    doc.add_heading("7.2. Transformación de TrackingTime", level=2)
    doc.add_paragraph("Archivo: processing/transformations.py → función transform_trackingtime()")
    doc.add_paragraph("Transformaciones aplicadas:")

    tt_transforms = [
        ("Eliminación de columnas innecesarias",
         "Se eliminan las columnas Service y Client, ya que no se utilizan en el reporte final."),
        ("Reemplazo de proyectos vacíos",
         "Los registros de tiempo que no tienen un proyecto asignado (valor vacío o nulo) "
         "se les asigna automáticamente el proyecto \"Administración\". Esto es porque ese tiempo "
         "se considera como trabajo administrativo general."),
        ("Normalización de nombres de proyecto",
         "Se corrigen inconsistencias en los nombres: \"Controles de Cambio y Errores\" → "
         "\"Controles de Cambios y Errores\" (agregar 's' faltante) y \"Integraciones Clientes\" → "
         "\"INTEGRACIONES\" (unificar nomenclatura con Azure)."),
        ("Filtrado de filas sin usuario",
         "Se eliminan las filas donde el campo User está vacío o es nulo."),
    ]
    for name, desc in tt_transforms:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(name + ": ")
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # 7.3 Tablas derivadas
    doc.add_heading("7.3. Construcción de Tablas Derivadas", level=2)
    doc.add_paragraph("Archivo: processing/transformations.py → función build_tables()")
    doc.add_paragraph(
        "A partir de los datos ya transformados y los archivos de referencia, se construyen las siguientes "
        "tablas intermedias:"
    )

    doc.add_heading("Tabla 1: total_horas_usuario_azure", level=3)
    doc.add_paragraph(
        "Agrupa los datos de Azure DevOps por Assigned To Clean y suma el campo Completed Work. "
        "Resultado: una fila por usuario con el total de horas registradas en Azure DevOps."
    )

    doc.add_heading("Tabla 2: total_horas_usuario_tracking", level=3)
    doc.add_paragraph(
        "Agrupa los datos de TrackingTime por User y suma el campo Hours. "
        "Resultado: una fila por usuario con el total de horas registradas en TrackingTime."
    )

    doc.add_heading("Tabla 3: total_horas_combinado", level=3)
    doc.add_paragraph(
        "Combina las tablas 1 y 2 (concatenación vertical) y vuelve a agrupar por Usuario "
        "sumando las horas totales de ambas fuentes. Luego se hace un LEFT JOIN con Equipo.csv "
        "usando el campo Nombre Azure como llave de cruce. El resultado final usa la columna "
        "NOMBRE de Equipo.csv como identificador del usuario (nombre completo formal)."
    )
    doc.add_paragraph(
        "Esta tabla sirve como referencia para calcular el porcentaje de asignación en la hoja de presentación."
    )

    doc.add_heading("Tabla 4: horas_azure (horas por proyecto — Azure)", level=3)
    doc.add_paragraph(
        "Agrupa los datos de Azure por Team Project + Assigned To Clean, sumando Completed Work. "
        "Se enriquece con dos LEFT JOINs:"
    )
    joins_az = [
        "Equipo.csv: Para obtener EMPRESA, NOMBRE completo y CÉDULA DE IDENTIDAD.",
        "CodigosProyectos.csv: Para obtener el Codigo Proyecto y Proyecto oficial "
        "(usando el campo Nombre Azure del catálogo como llave de cruce).",
    ]
    for j in joins_az:
        doc.add_paragraph(j, style='List Bullet 2')
    doc.add_paragraph("Se le agrega la columna Origen con valor fijo \"Azure\".")

    doc.add_heading("Tabla 5: horas_tracking (horas por proyecto — TrackingTime)", level=3)
    doc.add_paragraph(
        "Agrupa los datos de TrackingTime por Project + User, sumando Hours. "
        "Se enriquece con dos LEFT JOINs:"
    )
    joins_tr = [
        "Equipo.csv: Para obtener EMPRESA, NOMBRE completo y CÉDULA DE IDENTIDAD.",
        "CodigosProyectos.csv: Para obtener el Codigo Proyecto y Proyecto oficial "
        "(usando el campo Nombre Tracking del catálogo como llave de cruce).",
    ]
    for j in joins_tr:
        doc.add_paragraph(j, style='List Bullet 2')
    doc.add_paragraph("Se le agrega la columna Origen con valor fijo \"Tracking\".")

    # 7.4 Tabla Presentación
    doc.add_heading("7.4. Tabla de Presentación Final", level=2)
    doc.add_paragraph("Archivo: processing/transformations.py → función build_presentation_table()")
    doc.add_paragraph(
        "Esta es la tabla principal que se muestra en la primera pestaña del reporte Excel. "
        "Se construye así:"
    )
    steps_pres = [
        "Se concatenan verticalmente las tablas horas_azure y horas_tracking (unión de filas).",
        "Se hace un LEFT JOIN con total_horas_combinado para agregar la columna Total_Horas "
        "(total de horas del usuario en ambas fuentes).",
        "Se calcula una nueva columna \"% ASIGNADO A PROYECTOS\" = Suma / Total_Horas. "
        "Esto representa qué porcentaje del tiempo total de una persona fue dedicado a "
        "un proyecto específico.",
        "Se elimina la columna auxiliar Total_Horas del resultado final.",
    ]
    for i, step in enumerate(steps_pres, 1):
        doc.add_paragraph(f"{i}. {step}", style='List Number')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 8. INFORME EXCEL — ESTRUCTURA
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading("8. Informe Excel Final — Estructura del Archivo", level=1)
    doc.add_paragraph(
        "El archivo de salida se llama Reporte_YYYYMMDD_YYYYMMDD.xlsx y se genera en la carpeta output/. "
        "Contiene 4 pestañas (hojas) que se describen a continuación."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Formato visual: ")
    run.bold = True
    p.add_run("Todas las hojas tienen encabezados con fondo azul oscuro (#1F3864) y texto blanco en negrita. "
              "Las filas de datos alternan entre fondo blanco y gris claro (#F2F2F2) para facilitar la lectura. "
              "La fuente es Arial 10pt. Los anchos de columna se ajustan automáticamente al contenido.")

    doc.add_paragraph()

    # 8.1 PRESENTACION
    doc.add_heading("8.1. Pestaña: PRESENTACION", level=2)
    doc.add_paragraph(
        "Esta es la hoja principal del reporte. Muestra las horas trabajadas por persona, "
        "discriminadas por proyecto, con el porcentaje de tiempo dedicado a cada proyecto."
    )

    pres_cols = [
        ("EMPRESA", "Empresa a la que pertenece el colaborador (LOGIZTIK o NEWDATA). "
         "Proviene del cruce con Equipo.csv."),
        ("NOMBRE", "Nombre completo formal del colaborador. "
         "Proviene del cruce con Equipo.csv."),
        ("CÉDULA DE IDENTIDAD", "Número de cédula del colaborador. "
         "Proviene del cruce con Equipo.csv."),
        ("PROYECTO", "Nombre oficial del proyecto. "
         "Proviene del cruce con CodigosProyectos.csv."),
        ("CODIGO", "Código interno del proyecto (formato LAG-NNN-SIGLA). "
         "Proviene del cruce con CodigosProyectos.csv."),
        ("Suma", "Cantidad de horas trabajadas por esa persona en ese proyecto específico, "
         "provenientes de una sola fuente (Azure o TrackingTime). Formato: decimal con 1 cifra (ej: 42.5)."),
        ("Origen", "Fuente de datos de donde proviene el registro: \"Azure\" o \"Tracking\". "
         "Una persona puede tener registros de ambas fuentes para el mismo proyecto."),
        ("% ASIGNADO A PROYECTOS", "Porcentaje del tiempo total de la persona que fue dedicado "
         "a este proyecto. Se calcula como: Suma ÷ Total_Horas_del_usuario. "
         "Formato: porcentaje con 1 decimal (ej: 15.3%). La suma de todos los porcentajes "
         "de una persona (considerando todas sus filas) debería ser cercana al 100%."),
    ]
    add_styled_table(doc, ["Columna", "Descripción"], pres_cols, col_widths=[4.5, 13.5])

    doc.add_paragraph()

    # 8.2 CODIGOS PROYECTOS
    doc.add_heading("8.2. Pestaña: CODIGOS PROYECTOS", level=2)
    doc.add_paragraph(
        "Copia directa del archivo Documentos/CodigosProyectos.csv. Se incluye como referencia "
        "para que el lector del reporte pueda ver el catálogo completo de proyectos, "
        "sus descripciones y la distribución de costos por sucursal."
    )
    doc.add_paragraph(
        "Las columnas son idénticas a las descritas en la sección 4.1 de este documento. "
        "Los porcentajes de distribución (ECU, COL, USA, NL) se presentan como valores decimales "
        "(ej: 0.25 = 25%)."
    )

    # 8.3 DETALLE AZURE
    doc.add_heading("8.3. Pestaña: DETALLE - AZURE", level=2)
    doc.add_paragraph(
        "Contiene el detalle completo de los Work Items de Azure DevOps después de las "
        "transformaciones de limpieza (Fase 2-A). Un registro por cada Work Item."
    )

    azure_detail_cols = [
        ("ID", "Identificador numérico único del Work Item."),
        ("State", "Estado actual: New, Active, Closed, Resolved, etc."),
        ("Completed Work", "Horas registradas como completadas en este Work Item. Tipo numérico."),
        ("Team Project", "Nombre del proyecto en Azure DevOps."),
        ("Work Item Type", "Tipo de elemento: Task, Bug, User Story, Feature, etc."),
        ("Iteration Path", "Ruta del sprint o iteración (ej: \"AC Cliente 360\\Sprint 45\")."),
        ("Title", "Título descriptivo de la tarea."),
        ("Created Date", "Fecha de creación en formato M/D/YYYY h:mm:ss AM/PM."),
        ("Closed Date", "Fecha de cierre. Vacío si no se ha cerrado."),
        ("Changed Date", "Fecha de última modificación."),
        ("Assigned To Clean", "Nombre limpio del responsable (después de eliminar empresa y email)."),
    ]
    add_styled_table(doc, ["Columna", "Descripción"], azure_detail_cols, col_widths=[4, 14])

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Nota: ")
    run.bold = True
    p.add_run("La columna original \"Assigned To\" ha sido reemplazada por \"Assigned To Clean\" "
              "como resultado de la transformación descrita en la sección 7.1.")

    doc.add_page_break()

    # 8.4 DETALLE TRACKINGTIME
    doc.add_heading("8.4. Pestaña: DETALLE - TRACKINGTIME", level=2)
    doc.add_paragraph(
        "Contiene el detalle completo de los eventos de TrackingTime después de las "
        "transformaciones de limpieza (Fase 2-B). Un registro por cada evento de tiempo."
    )

    tt_detail_cols = [
        ("Project", "Nombre del proyecto en TrackingTime. Los registros sin proyecto "
         "se renombran a \"Administración\"."),
        ("User", "Nombre del usuario que registró el tiempo."),
        ("Hours", "Cantidad de horas del evento de tiempo. Tipo numérico (decimal con punto)."),
    ]
    add_styled_table(doc, ["Columna", "Descripción"], tt_detail_cols, col_widths=[4, 14])

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Nota: ")
    run.bold = True
    p.add_run("Las columnas originales Service y Client han sido eliminadas en la transformación "
              "descrita en la sección 7.2, ya que no aportan valor al reporte.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 9. FLUJO DE DATOS
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading("9. Flujo de Datos Completo (Resumen)", level=1)
    doc.add_paragraph(
        "A continuación se resume el flujo de datos desde la fuente hasta el reporte final:"
    )

    flow_steps = [
        ("ENTRADA – APIs externas",
         "Azure DevOps API → Work Items (ID, State, Assigned To, Completed Work, Team Project, etc.)\n"
         "TrackingTime API → Eventos de tiempo (Project, User, Hours, Service, Client)"),
        ("FASE 1 – Extracción",
         "Se descargan los datos crudos y se guardan como CSV intermedios en output/."),
        ("ENTRADA – Catálogos locales",
         "CodigosProyectos.csv → Mapeo de nombres de proyecto a códigos y sucursales.\n"
         "Equipo.csv → Mapeo de nombres de usuario a datos formales (nombre completo, empresa, cédula)."),
        ("FASE 2A – Limpieza",
         "Azure: Limpiar nombres (quitar empresa/email), filtrar vacíos.\n"
         "TrackingTime: Eliminar columnas innecesarias, normalizar nombres de proyecto, asignar 'Administración' a vacíos."),
        ("FASE 2B – Agregación",
         "Agrupar horas por usuario y por proyecto+usuario, desde cada fuente.\n"
         "Combinar ambas fuentes para obtener el total global por persona."),
        ("FASE 2C – Enriquecimiento",
         "Cruzar con Equipo.csv para obtener nombre formal, empresa y cédula.\n"
         "Cruzar con CodigosProyectos.csv para obtener código de proyecto y nombre oficial."),
        ("FASE 2D – Presentación",
         "Calcular % de asignación (horas en proyecto / total horas de la persona).\n"
         "Construir la tabla final y generar el archivo Excel con 4 pestañas."),
        ("SALIDA",
         "Reporte_YYYYMMDD_YYYYMMDD.xlsx con 4 pestañas:\n"
         "PRESENTACION | CODIGOS PROYECTOS | DETALLE - AZURE | DETALLE - TRACKINGTIME"),
    ]
    for title, desc in flow_steps:
        p = doc.add_paragraph()
        run = p.add_run(f"▶ {title}")
        run.bold = True
        run.font.color.rgb = RGBColor(31, 56, 100)
        for line in desc.split("\n"):
            doc.add_paragraph(line, style='List Bullet 2')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 10. CONFIGURACIÓN
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading("10. Configuración y Variables de Entorno", level=1)
    doc.add_paragraph(
        "Toda la configuración del programa se gestiona mediante un archivo .env en la raíz del proyecto. "
        "A continuación se describen las variables requeridas:"
    )

    env_vars = [
        ("TT_TOKEN", "Token de autenticación de TrackingTime (formato: Basic <base64>). "
         "Se genera como App Password desde la plataforma de TrackingTime."),
        ("TT_BASE_URL", "URL base de la API de TrackingTime. Valor por defecto: https://api.trackingtime.co/api/v4"),
        ("TT_DATE_FROM", "Fecha de inicio del rango del reporte (formato: YYYY-MM-DD). "
         "Define el inicio del período para la extracción de TrackingTime."),
        ("TT_DATE_TO", "Fecha fin del rango del reporte (formato: YYYY-MM-DD). "
         "Define el fin del período para la extracción de TrackingTime."),
        ("ADO_TOKEN", "Personal Access Token (PAT) de Azure DevOps (formato: Basic <base64>). "
         "Tiene fecha de vencimiento; debe renovarse periódicamente."),
        ("ADO_ORG", "Nombre de la organización en Azure DevOps (ej: lagcloud)."),
        ("ADO_PROJECT_ID", "ID del proyecto en Azure DevOps (GUID)."),
        ("ADO_QUERY_ID", "ID del query guardado en Azure DevOps que filtra los Work Items deseados (GUID). "
         "Los filtros de fecha y otras exclusiones se configuran directamente en el query de Azure DevOps."),
        ("ADO_BASE_URL", "URL base de Azure DevOps. Valor por defecto: https://dev.azure.com"),
    ]
    add_styled_table(doc, ["Variable", "Descripción"], env_vars, col_widths=[4, 14])

    doc.add_paragraph()

    doc.add_heading("Consideraciones importantes:", level=2)
    considerations = [
        "El rango de fechas para Azure DevOps NO se configura en el .env; se configura directamente "
        "en el query de Azure DevOps. Las variables TT_DATE_FROM y TT_DATE_TO solo controlan "
        "la extracción de TrackingTime y el sufijo de nombre de los archivos de salida.",
        "Los PAT de Azure DevOps tienen fecha de expiración. Si el programa falla al conectarse, "
        "verificar la vigencia del token.",
        "Si se agregan nuevos proyectos o personas al equipo, es obligatorio actualizar los archivos "
        "CodigosProyectos.csv y Equipo.csv en la carpeta Documentos/.",
        "Las dependencias del programa están listadas en requirements.txt y se instalan con: "
        "pip install -r requirements.txt",
    ]
    for c in considerations:
        doc.add_paragraph(c, style='List Bullet')

    # ── Guardar ──
    output_dir = "output"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "Informe_Descriptivo_Entregable.docx")
    doc.save(output_path)
    print(f"\n✅ Informe generado exitosamente: {output_path}")
    return output_path


if __name__ == "__main__":
    create_report()
